from flask import Flask, request, jsonify
import difflib
from docx import Document
from tempfile import TemporaryDirectory
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
import subprocess
import base64

app = Flask(__name__)

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    except ImportError:
        try:
            subprocess.run(['unoconv', '-f', 'pdf', '-o', os.path.dirname(pdf_path), docx_path], check=True)
        except:
            subprocess.run(['soffice', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path], check=True)

def get_words(text):
    """Split text into words for comparison."""
    return text.split()

def word_diff(old_words, new_words):
    """Generate a word-by-word diff using SequenceMatcher."""
    matcher = difflib.SequenceMatcher(None, old_words, new_words)
    return matcher.get_opcodes()

@app.route('/compare', methods=['POST'])
def compare_docs():
    try:
        files = request.get_json()
        if not isinstance(files, list) or len(files) < 2:
            return jsonify({"status": False, "message": "At least two files are required"}), 400

        with TemporaryDirectory() as temp_dir:
            # Decode and save both files
            def decode_and_save(file, idx):
                buffer = base64.b64decode(file["$content"])
                temp_path = os.path.join(temp_dir, f"file{idx}.docx")
                with open(temp_path, "wb") as f:
                    f.write(buffer)
                return temp_path

            path1 = decode_and_save(files[0], 1)
            path2 = decode_and_save(files[1], 2)

            # Convert to PDF
            pdf1 = os.path.join(temp_dir, "file1.pdf")
            pdf2 = os.path.join(temp_dir, "file2.pdf")
            convert_docx_to_pdf(path1, pdf1)
            convert_docx_to_pdf(path2, pdf2)

            # Extract all text from paragraphs
            doc1_text = "\n".join([para.text for para in Document(path1).paragraphs if para.text.strip()])
            doc2_text = "\n".join([para.text for para in Document(path2).paragraphs if para.text.strip()])

            # Split into words
            old_words = get_words(doc1_text)
            new_words = get_words(doc2_text)

            # Generate word-by-word diff
            opcodes = word_diff(old_words, new_words)

            # Generate PDF with highlighted word changes
            diff_pdf_path = os.path.join(temp_dir, "differences.pdf")
            c = canvas.Canvas(diff_pdf_path, pagesize=letter)
            width, height = letter
            y = height - 40

            c.setFont("Helvetica", 12)
            c.drawString(100, y, "Differences between documents (word-level):")
            y -= 20

            for op, i1, i2, j1, j2 in opcodes:
                if op == 'equal':
                    # Unchanged words: black
                    c.setFillColorRGB(0, 0, 0)
                    for word in old_words[i1:i2]:
                        c.drawString(60, y, word)
                        y -= 15
                elif op == 'delete':
                    # Deleted words: red
                    c.setFillColorRGB(0.8, 0, 0)
                    for word in old_words[i1:i2]:
                        c.drawString(60, y, f"[-{word}]")
                        y -= 15
                elif op == 'insert':
                    # Inserted words: green
                    c.setFillColorRGB(0, 0.6, 0)
                    for word in new_words[j1:j2]:
                        c.drawString(60, y, f"[+{word}]")
                        y -= 15
                elif op == 'replace':
                    # Replaced words: first deleted, then inserted
                    c.setFillColorRGB(0.8, 0, 0)
                    for word in old_words[i1:i2]:
                        c.drawString(60, y, f"[-{word}]")
                        y -= 15
                    c.setFillColorRGB(0, 0.6, 0)
                    for word in new_words[j1:j2]:
                        c.drawString(60, y, f"[+{word}]")
                        y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()

            # Read all PDFs and encode as base64
            def encode_pdf(pdf_path):
                with open(pdf_path, "rb") as f:
                    return base64.b64encode(f.read()).decode('utf-8')

            result = [
                {
                    "$content-type": "application/pdf",
                    "$content": encode_pdf(pdf1)
                },
                {
                    "$content-type": "application/pdf",
                    "$content": encode_pdf(pdf2)
                },
                {
                    "$content-type": "application/pdf",
                    "$content": encode_pdf(diff_pdf_path)
                }
            ]
            return jsonify(result)

    except Exception as e:
        return jsonify({"status": False, "message": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
